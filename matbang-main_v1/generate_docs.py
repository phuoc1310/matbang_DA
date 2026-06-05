import psycopg2
import docx
from docx.shared import Inches

def get_db_schema():
    conn = psycopg2.connect(
        host="localhost",
        database="mat_bang",
        user="postgres",
        password="123456",
        port=5432
    )
    cur = conn.cursor()
    
    cur.execute("""
        SELECT table_name 
        FROM information_schema.tables 
        WHERE table_schema = 'public' 
        ORDER BY table_name;
    """)
    tables = [row[0] for row in cur.fetchall()]
    
    schema = {}
    for table in tables:
        cur.execute(f"""
            SELECT column_name, data_type, character_maximum_length 
            FROM information_schema.columns 
            WHERE table_schema = 'public' AND table_name = '{table}' 
            ORDER BY ordinal_position;
        """)
        columns = cur.fetchall()
        schema[table] = columns
        
    cur.close()
    conn.close()
    return schema

def create_markdown(schema, md_path):
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write("# Thiết kế CSDL (Đã cập nhật từ PostgreSQL)\n\n")
        
        for i, (table, columns) in enumerate(schema.items(), 1):
            f.write(f"## {i}. Bảng {table.capitalize()}\n\n")
            f.write(f"Bảng {table.capitalize()} lưu trữ thông tin liên quan đến {table}.\n\n")
            f.write("| Tên cột | Kiểu dữ liệu | Độ dài tối đa |\n")
            f.write("|---|---|---|\n")
            for col in columns:
                col_name, data_type, max_len = col
                max_len_str = str(max_len) if max_len is not None else "-"
                f.write(f"| {col_name} | {data_type} | {max_len_str} |\n")
            f.write("\n")

def create_docx(schema, docx_path):
    doc = docx.Document()
    doc.add_heading("Thiết kế CSDL (Đã cập nhật từ PostgreSQL)", 0)
    
    for i, (table, columns) in enumerate(schema.items(), 1):
        doc.add_heading(f"{i}. Bảng {table.capitalize()}", level=2)
        doc.add_paragraph(f"Bảng {table.capitalize()} lưu trữ thông tin liên quan đến {table}.")
        
        table_doc = doc.add_table(rows=1, cols=3)
        table_doc.style = 'Table Grid'
        hdr_cells = table_doc.rows[0].cells
        hdr_cells[0].text = 'Tên cột'
        hdr_cells[1].text = 'Kiểu dữ liệu'
        hdr_cells[2].text = 'Độ dài tối đa'
        
        for col in columns:
            row_cells = table_doc.add_row().cells
            col_name, data_type, max_len = col
            max_len_str = str(max_len) if max_len is not None else "-"
            row_cells[0].text = str(col_name)
            row_cells[1].text = str(data_type)
            row_cells[2].text = str(max_len_str)
            
        doc.add_paragraph()
        
    doc.save(docx_path)

if __name__ == '__main__':
    schema = get_db_schema()
    create_markdown(schema, r'c:\Users\PHUOC\OneDrive\Documents\DA\mat_bang\thiet_ke_csdl_final.md')
    create_docx(schema, r'c:\Users\PHUOC\OneDrive\Documents\DA\mat_bang\thiet_ke_csdl_final.docx')
