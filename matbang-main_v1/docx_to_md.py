import docx
import sys

def docx_to_md(docx_path, md_path):
    doc = docx.Document(docx_path)
    with open(md_path, 'w', encoding='utf-8') as f:
        for para in doc.paragraphs:
            if para.text.strip():
                f.write(para.text + '\n\n')
        
        for table in doc.tables:
            for i, row in enumerate(table.rows):
                row_data = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                f.write('| ' + ' | '.join(row_data) + ' |\n')
                if i == 0:
                    f.write('|' + '|'.join(['---'] * len(row.cells)) + '|\n')
            f.write('\n\n')

if __name__ == '__main__':
    docx_to_md(sys.argv[1], sys.argv[2])
