import docx

def md_to_docx(md_path, docx_path):
    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()

        doc = docx.Document()
        
        in_code_block = False
        code_text = []

        for line in content.split('\n'):
            if line.startswith('```'):
                if in_code_block:
                    in_code_block = False
                    p = doc.add_paragraph('\n'.join(code_text))
                    # Optional: style it as code
                    # p.style = 'Macro Text'
                    code_text = []
                else:
                    in_code_block = True
                    code_text = []
                continue
            
            if in_code_block:
                code_text.append(line)
                continue

            if line.startswith('# '):
                doc.add_heading(line[2:].strip(), 0)
            elif line.startswith('## '):
                doc.add_heading(line[3:].strip(), 1)
            elif line.startswith('### '):
                doc.add_heading(line[4:].strip(), 2)
            elif line.startswith('#### '):
                doc.add_heading(line[5:].strip(), 3)
            elif line.startswith('---'):
                pass # skip hr
            elif line.strip() == '':
                pass # doc.add_paragraph()
            else:
                p = doc.add_paragraph()
                parts = line.split('**')
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1:
                        run.bold = True


        doc.save(docx_path)
        print(f"Successfully generated {docx_path}")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == '__main__':
    md_to_docx(
        r'c:\Users\PHUOC\OneDrive\Documents\DA\mat_bang\matbang-main_v1\use_case_diagrams.md',
        r'c:\Users\PHUOC\OneDrive\Documents\DA\mat_bang\matbang-main_v1\use_case_diagrams.docx'
    )
